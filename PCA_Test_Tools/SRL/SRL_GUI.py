import pandas as pd
# import numpy as np
# import os
# from os import path
import pickle
# from collections import defaultdict
# import plotly.express as px
# import plotly.subplots as sp
# import plotly.graph_objects as go
# import scipy.stats as stats
# import open3d as o3d
# import logging
# from scipy.stats import binom
# import glob
import gzip
# from pathlib import Path
# import pathlib
# import session_info
# from ecal.measurement.measurement import Measurement
# from selda.pc2_processor import PC2Processor
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import warnings
warnings.filterwarnings('ignore')

import sys
sys.path.append('../libraries/')
import ingest
import Calculations


#tkinter library is used to create the GUI
import tkinter as tk
from tkinter import filedialog #(needed for path selection button)
from tkinter import ttk #needed for the combobox (drop down menu)
from tkinter.messagebox import askyesno #needed for Exit button

#------------------------------------------------------------------------------------------------------------------------------
#function needed to create tabs
def tab_changed(event):
    selected_tab = tab_control.index(tab_control.select())
    print(f"Selected Tab: {selected_tab}")

#------------------------------------------------------------------------------------------------------------------------------
#function needed to create Exit button
def exit_app():
    if askyesno("Exit", "Are you sure you want to exit?"):
        root.destroy()

#------------------------------------------------------------------------------------------------------------------------------
#function to select a file path
def select_file(entry):
    file_path = filedialog.askopenfilename()
    entry.delete(0, tk.END)  # Clear the current entry
    entry.insert(0, file_path)  # Insert the selected file path into the entry field

#------------------------------------------------------------------------------------------------------------------------------
#function to select a folder path
def select_HDF5_path(entry):
    folder_path = filedialog.askdirectory()
    if folder_path:  # Check if folder_path is not empty
        entry.delete(0, tk.END)
        entry.insert(0, folder_path)
        set_sensor_button.config(state=tk.NORMAL)  # Enable the button
        log_text.insert(tk.END, f"Click Set Sensor\n\n")
    # entry.delete(0, tk.END)  # Clear the current entry
    # entry.insert(0, folder_path)  # Insert the selected folder path into the entry field

#------------------------------------------------------------------------------------------------------------------------------
#function to select a folder path
def select_pkl_path(entry):
    folder_path = filedialog.askdirectory()
    if folder_path:  # Check if folder_path is not empty
        entry.delete(0, tk.END)
        entry.insert(0, folder_path)
        save_to_csv_button.config(state=tk.NORMAL)  # Enable the button
    # entry.delete(0, tk.END)  # Clear the current entry
    # entry.insert(0, folder_path)  # Insert the selected folder path into the entry field
#------------------------------------------------------------------------------------------------------------------------------
# function to handle file upload button click
def upload_file():
    file_path = filedialog.askopenfilename(filetypes=[("Pickle files", "*.pkl")])
    if file_path:
        pkl_file_name_label.config(text="Selected file: " + file_path)

#------------------------------------------------------------------------------------------------------------------------------
# function to handle button check
# def on_check_click():
#     sample_type = option_combobox.get()  #get the text from the option selection of the sensor
#     directory = input_entry1.get()  # Get text from input entry 1
#     pkl_directory = input_entry2.get()  # Get text from input entry 2
#     result = f"Your sensor is: {sample_type}\n You entered:\nRecording path: {directory}\nOutput pkl path: {pkl_directory}"
#     result_label.config(text=result)


#------------------------------------------------------------------------------------------------------------------------------
# function to handle button check

def on_set_sensor_click():
    sample_type = option_combobox.get()  #get the text from the option selection of the sensor
    directory = input_entry1.get()  # Get text from input entry 1
    global sensor
    sensor = ingest.sensor(sample_type,directory)
    print(sensor.sensor_type)
    log_text.insert(tk.END, f"Selected option: {sample_type}\n")
    log_text.insert(tk.END, f"Selected directory: {directory}\n\n")
    log_text.insert(tk.END, f"Click Ingest Data\n\n")
    ingest_button.config(state=tk.NORMAL)
    status = f"Sensor Set to {sensor.sensor_type}"
    result_label.config(text=status)
    




#------------------------------------------------------------------------------------------------------------------------------
# function to handle button ingest
def on_ingest_click():
    ingest.data_read_new(sensor)
    log_text.insert(tk.END, f"Row Data Frame:\n")
    log_text.insert(tk.END, f"{sensor.df}\n\n")

    ingest.row_col_fix_new(sensor)
    log_text.insert(tk.END, f"Data Frame after row/col fix:\n")
    log_text.insert(tk.END, f"{sensor.df}\n\n")

    ingest.data_validation_new(sensor, frame_number = 5)
    log_text.insert(tk.END, f"Valid Returns:\n")
    log_text.insert(tk.END, f"{sensor.df_valid_returns}\n\n")

    log_text.insert(tk.END, f"Ingest Complete!\n\n")

    log_text.insert(tk.END, f"If you wish to save DF to csv to view, choose path for storing files\n\n")

#------------------------------------------------------------------------------------------------------------------------------
# function to save processed dataframe to csv 
def on_save_to_csv():
    pkl_directory = input_entry2.get()  # Get text from input entry 2
    csvpath=pkl_directory+"/processed_dataframe.csv"
    sensor.df_valid_returns.to_csv(csvpath, index=False)
    log_text.insert(tk.END, f"dataframe saved to csv\n\n")

#------------------------------------------------------------------------------------------------------------------------------
# function to handle button save pkl
def on_save_pkl_click():
    directory = input_entry1.get()  # Get text from input entry 1
    pkl_directory = input_entry2.get()  # Get text from input entry 2
    #global df_import
    pkl_file = pkl_directory + "All_frames.pkl"
    df_import.to_pickle(pkl_file)

#------------------------------------------------------------------------------------------------------------------------------
# function to handle button compress pkl
def on_compress_pkl_click():
    directory = input_entry1.get()  # Get text from input entry 1
    pkl_directory = input_entry2.get()  # Get text from input entry 2
    pkl_file = pkl_directory + "All_frames.pkl"
    with gzip.open(directory + 'compressed_pkl.gz', 'wb') as f: pickle.dump(pkl_file, f)

#------------------------------------------------------------------------------------------------------------------------------
# function to select target
def on_target_select():
    Calculations.target_selection_new(sensor)


#------------------------------------------------------------------------------------------------------------------------------
# function to handle button PFA
def on_PFA_click():
    global pfa
    pfa = Calculations.pfa_calculation(sensor)

    PFA_label.config(text=f"PFA = {pfa}")
    PFA_label.grid(row=4,column=0)
    return pfa

#------------------------------------------------------------------------------------------------------------------------------
# function to handle button PD
def on_PD_click(pd):
    PD_label.config(text="PD = {pd}")

#------------------------------------------------------------------------------------------------------------------------------
# function to handle button Ground Plane
def on_GP_click():
    pass

#------------------------------------------------------------------------------------------------------------------------------
# function to handle button PPS
def on_PPS_click():
    pass

#------------------------------------------------------------------------------------------------------------------------------
# function to handle button Target Size
def on_TS_click():
    pass

# #------------------------------------------------------------------------------------------------------------------------------
# Function to handle checkbox state change (optimized)
def on_checkbox_click():
    # print("PFA State:", checkbox_vars["PFA"].get())
    # print("PD State:", checkbox_vars["PD"].get())
    # print("GT State:", checkbox_vars["GT"].get())
    # print("PPS State:", checkbox_vars["PPS"].get())
    # print("TS State:", checkbox_vars["TS"].get())
    pfa_state = checkbox_vars["PFA"].get()
    pd_state = checkbox_vars["PD"].get()
    gt_state = checkbox_vars["GT"].get()
    pps_state = checkbox_vars["PPS"].get()
    ts_state = checkbox_vars["TS"].get()
    log_text.insert(tk.END, f"PFA: {pfa_state}\tPD: {pd_state}\tGT: {gt_state}\tPPS: {pps_state}\tTS: {ts_state}\n\n")


# # function to handle the PFA checkbox state change
# def on_PFA_checkbox_click():
#     checkbox_value = PFA_state.get()
#     print("PFA state:", checkbox_value)

# #------------------------------------------------------------------------------------------------------------------------------
# # function to handle the PD checkbox state change
# def on_PD_checkbox_click():
#     checkbox_value = PD_state.get()
#     print("PD state:", checkbox_value)

# #------------------------------------------------------------------------------------------------------------------------------
# # function to handle the GT checkbox state change
# def on_GT_checkbox_click():
#     checkbox_value = GT_state.get()
#     print("GT state:", checkbox_value)

# #------------------------------------------------------------------------------------------------------------------------------
# # function to handle the PPS checkbox state change
# def on_PPS_checkbox_click():
#     checkbox_value = PPS_state.get()
#     print("PPS state:", checkbox_value)

# #------------------------------------------------------------------------------------------------------------------------------
# # function to handle the TS checkbox state change
# def on_TS_checkbox_click():
#     checkbox_value = TS_state.get()
#     print("TS state:", checkbox_value)

#------------------------------------------------------------------------------------------------------------------------------
# function to handle button calculate
def on_calculate_click():
    pass

#------------------------------------------------------------------------------------------------------------------------------
# Function to handle export button click
def export_to_word():
    # Create a new Word document
    doc = Document()
    sample_type = option_combobox.get()  #get the text from the option selection of the sensor
    # Add content to the Word document with specific formatt
    # ing
    title = doc.add_heading(f'{sample_type} KPI TEST Report\n', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the title
    title.bold = True  # Make the title bold

    PFA_header = doc.add_heading(f'Probability of False Alarm\n', level=2)
    PFA_header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Center-align the title
    PFA_header.bold = True  # Make the title bold

    # Add a paragraph with custom font size
    paragraph = doc.add_paragraph('PFA =\n')
    #paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justify-align the paragraph
    paragraph.add_run('\n')  # Add a line break
    paragraph.add_run('This paragraph has custom font size.').font.size = Pt(12)

    # Save the Word document
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    if file_path:
        doc.save(file_path)

        # Show message to the user
        tk.messagebox.showinfo("Success", "Report exported successfully.")

#------------------------------------------------------------------------------------------------------------------------------
#function to resize grid dynamically
def on_resize(event):
    # Get the new size of the main window
    width = event.width
    height = event.height
    
    # Adjust the size of the log_text widget
    log_text.config(width=int(width/20), height=int(height/100))


#------------------------------------------------------------------------------------------------------------------------------
# Create the main application window
root = tk.Tk()
root.title("SRL Analysis Tool")
#root.geometry("800x600")  # Width x Height
root.geometry("") #dynamic size of window

# Create a style for the Notebook widget
style = ttk.Style()
style.theme_create('OrangeTabs', parent='alt', settings={
    "TNotebook": {"configure": {"tabmargins": [2, 5, 2, 0]}},
    "TNotebook.Tab": {
        "configure": {"padding": [10, 5], "background": "orange"},
        "map": {"background": [("selected", "orange")]}
    }
})
style.theme_use('OrangeTabs')

# Create a Notebook widget
tab_control = ttk.Notebook(root)
tab_control.pack(fill='both', expand=True)

# Create and add tabs
tab1 = ttk.Frame(tab_control)
tab_control.add(tab1, text='Ingest')

tab2 = ttk.Frame(tab_control)
tab_control.add(tab2, text='Analysis')

tab3 = ttk.Frame(tab_control)
tab_control.add(tab3, text='Report')

    
#------------------------------------------------------------------------------------------------------------------------------
#------------------TAB1--------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------

#------------------------------------------------------------------------------------------------------------------------------
#create introduction text
intro_label = tk.Label(tab1, text="Follow the instructions in the interactive window", anchor=tk.W, justify=tk.LEFT,font=("Helvetica", 10))
intro_label.grid(row=0, column=0, columnspan=3, padx=5, pady=5,sticky="nsew")

#------------------------------------------------------------------------------------------------------------------------------
#create option selection box
option_label = tk.Label(tab1, text="Select a sensor:",font=("Helvetica", 10, "bold"))
option_label.grid(row=1, column=0, padx=5, pady=5)

options = ["voyant", "scantinel", "silc", "HRL", "SRL","Hessai"]  # List of options

option_combobox = ttk.Combobox(tab1, values=options, width=10)
option_combobox.set("SRL")  # Value of the default option
option_combobox.grid(row=1, column=1, columnspan=1, padx=5, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create input label and path selection input button for HDF5 path
input_label1 = tk.Label(tab1, text="Path for HDF5:",font=("Helvetica", 10, "bold"))
input_label1.grid(row=2, column=0, padx=5, pady=5)

input_entry1 = tk.Entry(tab1, width=70)
input_entry1.grid(row=2, column=1, padx=5, pady=5)

file_button1 = tk.Button(tab1, text="Select Folder", command=lambda: select_HDF5_path(input_entry1))
file_button1.grid(row=2, column=2, padx=5, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create set_sensor button
set_sensor_button = tk.Button(tab1, text="Set Sensor", command=on_set_sensor_click, width=10,state=tk.DISABLED)
# set_sensor_button = tk.Button(tab1, text="Set Sensor", command=on_set_sensor_click, width=10)
set_sensor_button.grid(row=1, column=2, padx=5, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create input label and path selection input button for .pkl files storage
input_label2 = tk.Label(tab1, text="Path for storing files:",font=("Helvetica", 10, "bold"))
input_label2.grid(row=3, column=0, padx=5, pady=5)

input_entry2 = tk.Entry(tab1, width=70)
input_entry2.grid(row=3, column=1, padx=5, pady=5)

file_button2 = tk.Button(tab1, text="Select Folder", command=lambda: select_pkl_path(input_entry2))
file_button2.grid(row=3, column=2, padx=5, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create check button
# check_button = tk.Button(tab1, text="Check", command=on_check_click, width=10)
# check_button.grid(row=4, column=1, padx=5, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create ingest button
ingest_button = tk.Button(tab1, text="Ingest Data", command=on_ingest_click, width=10, state=tk.DISABLED)
ingest_button.grid(row=4, column=0, padx=5, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create save to csv button
save_to_csv_button = tk.Button(tab1, text="Save DF (csv)", command=on_save_to_csv, width=10, state=tk.DISABLED)
save_to_csv_button.grid(row=5, column=0, padx=5, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create save pkl button
save_pkl_button = tk.Button(tab1, text="save .pkl", command=on_save_pkl_click, width=10)
save_pkl_button.grid(row=5, column=1, padx=5, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create compress pkl button
compress_pkl_button = tk.Button(tab1, text="compress .pkl", command=on_compress_pkl_click, width=10)
compress_pkl_button.grid(row=5, column=2, padx=5, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create ending text
ingest_completion_label = tk.Label(tab1, text="")
ingest_completion_label.grid(row=6, column=0, columnspan=3, padx=5, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create ending text
result_label = tk.Label(tab1, text="")
result_label.grid(row=7, column=0, columnspan=3, padx=5, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
#------------------TAB2--------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------


#------------------------------------------------------------------------------------------------------------------------------
# Create a button to select a target from point cloud
select_target_button = tk.Button(tab2, text="Select Target", command=on_target_select)
select_target_button.grid(row=2, column=1, padx=10, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# Create a button to upload a .pkl file
pkl_upload_button = tk.Button(tab2, text="Upload .pkl File", command=upload_file)
pkl_upload_button.grid(row=2, column=2, padx=10, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create PFA button
PFA_button = tk.Button(tab2, text="PFA", command=on_PFA_click, width=15, font=("Helvetica", 10, "bold"))
PFA_button.grid(row=3, column=0, padx=10, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create PD button
pd = 0
PD_button = tk.Button(tab2, text="PD", command=lambda: on_PD_click(pd), width=15, font=("Helvetica", 10, "bold"))
PD_button.grid(row=3, column=1, padx=10, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create Ground Plane button
GP_button = tk.Button(tab2, text="Ground Plane", command=on_GP_click, width=15, font=("Helvetica", 10, "bold"))
GP_button.grid(row=3, column=2, padx=10, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create Points per Second button
PPS_button = tk.Button(tab2, text="Points per Second", command=on_PPS_click, width=15, font=("Helvetica", 10, "bold"))
PPS_button.grid(row=3, column=3, padx=10, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create target size button
TS_button = tk.Button(tab2, text="Target Size", command=on_TS_click, width=15, font=("Helvetica", 10, "bold"))
TS_button.grid(row=3, column=4, padx=10, pady=5)





# #------------------------------------------------------------------------------------------------------------------------------
# # create PFA text
# PFA_label = tk.Label(tab2, text="")
# PFA_label.grid(row=4, column=0, padx=10, pady=5)

# #------------------------------------------------------------------------------------------------------------------------------
# # create PD text
# PD_label = tk.Label(tab2, text="")
# PD_label.grid(row=4, column=1, padx=10, pady=5)

# #------------------------------------------------------------------------------------------------------------------------------
# # create GP text
# GP_label = tk.Label(tab2, text="")
# GP_label.grid(row=4, column=2, padx=10, pady=5)

# #------------------------------------------------------------------------------------------------------------------------------
# # create PPS text
# PPS_label = tk.Label(tab2, text="")
# PPS_label.grid(row=4, column=3, padx=10, pady=5)

# #------------------------------------------------------------------------------------------------------------------------------
# # create TS text
# TS_label = tk.Label(tab2, text="")
# TS_label.grid(row=4, column=4, padx=10, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# Define checkbox labels
checkbox_labels = ["PFA", "PD", "GT", "PPS", "TS"]

# Create a dictionary to store BooleanVar variables
checkbox_vars = {}

# Create BooleanVar variables and checkboxes (optimized)
for i, label in enumerate(checkbox_labels):
    checkbox_vars[label] = tk.BooleanVar()
    checkbox = tk.Checkbutton(tab3, text=label, variable=checkbox_vars[label], command=on_checkbox_click)
    checkbox.grid(row=1, column=i, padx=10, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# create Calculate button
calculate_button = tk.Button(tab2, text="Calculate", command=on_calculate_click, width=15, font=("Helvetica", 10, "bold"))
calculate_button.grid(row=6, column=2, padx=10, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
#------------------TAB3--------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------

#------------------------------------------------------------------------------------------------------------------------------
# Create a button to export results to Word
export_button = tk.Button(tab3, text="Export Report", command=export_to_word)
export_button.grid(row=2, column=0, padx=10, pady=10)

#------------------------------------------------------------------------------------------------------------------------------
# Bind event handler for tab change
tab_control.bind("<<NotebookTabChanged>>", tab_changed)

#------------------------------------------------------------------------------------------------------------------------------
# Notebook widget
tab_control.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")

#------------------ROOT--------------------------------------------------------------------------------------------------------

#------------------------------------------------------------------------------------------------------------------------------
# Create a label to display the selected file name
pkl_file_name_label = tk.Label(root, text="")
pkl_file_name_label.grid(row=5, column=0, padx=5, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# Create a label for the log window
log_label = tk.Label(root, text="Interactive Window", font=("Helvetica", 12, "bold"))
log_label.grid(row=6, column=0, padx=5, pady=5)

#------------------------------------------------------------------------------------------------------------------------------
# Create a log window to display the logs
log_text = tk.Text(root, width=60, height=60, bg="black", fg="white")
#log_text.grid(row=7, column=0, columnspan=1, padx=10, pady=5, sticky="w")
log_text.grid(row=7, column=0, padx=15, pady=5, sticky="nsew")
log_text.insert(tk.END, f"Start by selecting a sensor and a path for HDF5 recordings\n\n")

#------------------------------------------------------------------------------------------------------------------------------
# Add an exit button
exit_button = tk.Button(root, text="Exit", command=exit_app, width=10, font=("Helvetica", 10, "bold"))
exit_button.grid(row=8, column=0, columnspan=3, padx=10, pady=20)

#------------------------------------------------------------------------------------------------------------------------------
# Configure row and column to resize proportionally
root.grid_rowconfigure(0, weight=1)
root.grid_columnconfigure(0, weight=1)

# Bind the resize event of the main window to the on_resize function
root.bind("<Configure>", on_resize)

# # Bind the resize event of the tab's frame to the on_resize function
# tab1.bind("<Configure>", on_resize)
# tab2.bind("<Configure>", on_resize)
# tab3.bind("<Configure>", on_resize)

#------------------------------------------------------------------------------------------------------------------------------
# Start the main event loop
root.mainloop()